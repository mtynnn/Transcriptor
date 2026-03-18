from faster_whisper import WhisperModel
import os
from docx import Document

class AudioTranscriber:
    def __init__(self, model_size="base", compute_type="int8", device="auto"):
        import ctranslate2
        self.model_size = model_size
        self.actual_device = "cpu"
        """
        Inicializa Faster-Whisper. Detecta automáticamente si hay GPU (CUDA).
        Si CUDA falla (ej: cublas64_12.dll no encontrada), cae a CPU automáticamente.
        """
        if device == "auto":
            # ctranslate2 es el motor interno de faster-whisper. Usarlo no falla en PyInstaller.
            if ctranslate2.get_cuda_device_count() > 0:
                print(f"[GPU] ¡NVIDIA (o CUDA) detectada! Usando aceleración por hardware.")
                device = "cuda"
                # En Windows, float16 es lo más rápido y estable para GPUs NVIDIA
                if compute_type == "int8":
                    compute_type = "float16"
            else:
                print("[CPU] No se detectó GPU CUDA compatible. Usando procesador.")
                device = "cpu"
                compute_type = "int8"

        print(f"Inicializando motor con modelo {model_size} en {device} ({compute_type})...")
        try:
            self.model = WhisperModel(model_size, device=device, compute_type=compute_type)
            self.actual_device = device
        except Exception as e:
            if device == "cuda":
                print(f"[FALLBACK] Error al iniciar CUDA: {e}")
                print("[FALLBACK] Reintentando con CPU. La transcripción funcionará pero será más lenta.")
                device = "cpu"
                compute_type = "int8"
                self.model = WhisperModel(model_size, device=device, compute_type=compute_type)
                self.actual_device = "cpu"
            else:
                raise
        print(f"Modelo inicializado correctamente en {self.actual_device}.")

    def _run_transcribe(self, audio_path: str, initial_prompt: str, device: str):
        """Ejecuta la transcripción en el device dado."""
        return self.model.transcribe(
            audio_path,
            initial_prompt=initial_prompt,
            vad_filter=True,
            vad_parameters=dict(min_silence_duration_ms=500),
            task="transcribe"
        )

    def transcribe(self, audio_path: str, initial_prompt: str = None, progress_callback=None) -> list:
        """
        Transcribe el audio y reporta el progreso real a través de un callback.
        Si la GPU se queda sin memoria durante la transcripción, cae a CPU automáticamente.
        """
        if not os.path.exists(audio_path):
            raise FileNotFoundError(f"Archivo de audio no encontrado: {audio_path}")

        print(f"Empezando transcripción. Prompt inicial: {initial_prompt[:50]}..." if initial_prompt else "Empezando transcripción sin prompt.")

        try:
            segments, info = self._run_transcribe(audio_path, initial_prompt, self.actual_device)
        except Exception as e:
            error_str = str(e).lower()
            is_oom = any(k in error_str for k in ("out of memory", "cuda error", "cublas", "cudnn", "failed to allocate"))
            if is_oom and self.actual_device == "cuda":
                print(f"[FALLBACK] Error de GPU durante transcripción: {e}")
                print("[FALLBACK] Recargando modelo en CPU para este archivo...")
                self.model = WhisperModel(self.model_size, device="cpu", compute_type="int8")
                self.actual_device = "cpu"
                segments, info = self._run_transcribe(audio_path, initial_prompt, "cpu")
            else:
                raise

        total_duration = info.duration
        print(f"Idioma detectado: {info.language} ({info.language_probability * 100:.2f}%) | Duración: {total_duration:.2f}s")

        segment_list = []
        for segment in segments:
            segment_list.append(segment)
            if progress_callback and total_duration > 0:
                # Calcular progreso basado en la posición actual del audio
                current_percent = int((segment.end / total_duration) * 100)
                progress_callback(min(current_percent, 99)) # Reservamos el 100 para el final

        if progress_callback:
            progress_callback(100)

        return segment_list

    def format_to_fluid_paragraphs(self, segments: list) -> str:
        """
        Convierte la lista de segmentos en un string continuo y legible.
        Aplica reglas de puntuación heurística basadas en los 'gaps' de audio.
        """
        if not segments:
            return ""

        processed_text = ""
        for i, segment in enumerate(segments):
            text = segment.text.strip()
            if not text:
                continue

            if i > 0:
                prev_segment = segments[i-1]
                gap = segment.start - prev_segment.end
                
                # Reglas de Puntuación Heurística:
                # 1. Salto de párrafo (Pausa larga > 1.8 seg o punto previo + pausa > 0.6 seg)
                ends_with_terminal = processed_text.strip()[-1] in [".", "?", "!"] if (processed_text and processed_text.strip()) else False
                
                if gap > 1.8 or (ends_with_terminal and gap > 0.6):
                    processed_text += "\n\n" + text.capitalize()
                # 2. Coma automática (Pausa breve entre 0.3 y 0.7 seg sin puntuación previa)
                elif 0.3 < gap < 0.7 and not ends_with_terminal:
                    processed_text += ", " + text.lower()
                # 3. Espacio simple (Pausa muy corta)
                else:
                    if ends_with_terminal:
                        processed_text += " " + text.capitalize()
                    else:
                        processed_text += " " + text
            else:
                processed_text = text.capitalize()

        return processed_text

    def save_as_txt(self, text: str, output_path: str):
        """Guarda el texto en un archivo plano."""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        return output_path

    def save_as_docx(self, text: str, output_path: str):
        """Guarda el texto en un archivo de Word formateado."""
        doc = Document()
        doc.add_heading('Transcripción Veterinaria Automática', 0)
        
        # Dividir por párrafos (detectados por \n\n)
        paragraphs = text.split("\n\n")
        for p_text in paragraphs:
            if p_text.strip():
                doc.add_paragraph(p_text.strip())
        
        doc.save(output_path)
        return output_path
